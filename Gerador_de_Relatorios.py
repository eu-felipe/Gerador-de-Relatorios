import pandas as pd
import os
import sys
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import win32com.client

# Variável global para controlar o cancelamento
cancelar_processo = False

# Função para garantir que o caminho do recurso funcione no executável
def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Função para sanitizar o nome do arquivo
def sanitizar_nome_arquivo(nome):
    # Remove caracteres inválidos para nomes de arquivos
    return re.sub(r'[\\/*?:"<>|]', '_', nome)

# Constantes para cores, fontes e mensagens
COR_TEXTO = RGBColor(255, 255, 255)  # Branco
COR_DESTAQUE = WD_COLOR_INDEX.DARK_YELLOW  # Destaque amarelo escuro
FONTE = "IBM Plex Sans"
TAMANHO_FONTE_TITULO = Pt(14)
MENSAGEM_SUCESSO = "Arquivos PDF gerados com sucesso!"
MENSAGEM_ERRO = "Ocorreu um erro ao gerar os PDFs: {}"
CAMINHO_LOGO = resource_path("Imagem2.png")  # Caminho da logo (usando caminho relativo)

def selecionar_arquivo():
    caminho = filedialog.askopenfilename(filetypes=[("Arquivos CSV", "*.csv")])
    entry_arquivo.delete(0, tk.END)
    entry_arquivo.insert(0, caminho)

def selecionar_pasta():
    pasta = filedialog.askdirectory()
    entry_pasta.delete(0, tk.END)
    entry_pasta.insert(0, pasta)

def converter_para_pdf(docx_file, pdf_file):
    try:
        if not os.path.exists(docx_file):
            messagebox.showerror("Erro", f"Arquivo DOCX não encontrado: {docx_file}")
            return
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_file)
        doc.SaveAs(pdf_file, FileFormat=17)
        doc.Close()
        word.Quit()
        
        # Remove o DOCX apenas após a conversão bem-sucedida
        if os.path.exists(docx_file):
            os.remove(docx_file)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao converter {docx_file} para PDF: {e}")
        raise

def cancelar():
    global cancelar_processo
    cancelar_processo = True
    messagebox.showinfo("Cancelado", "Processo de geração de PDFs cancelado.")

def gerar_pdfs():
    global cancelar_processo
    cancelar_processo = False  # Reinicia a variável de cancelamento

    try:
        arquivo_csv = entry_arquivo.get()
        pasta_saida = entry_pasta.get()
        
        if not arquivo_csv or not pasta_saida:
            messagebox.showwarning("Aviso", "Por favor, selecione um arquivo CSV e uma pasta de saída.")
            return
        
        if not os.path.isfile(arquivo_csv):
            messagebox.showerror("Erro", "O arquivo CSV selecionado não existe.")
            return
        
        if not os.access(pasta_saida, os.W_OK):
            messagebox.showerror("Erro", "A pasta de saída não tem permissão de escrita.")
            return
        
        df = pd.read_csv(arquivo_csv)
        total_linhas = len(df)
        
        os.makedirs(pasta_saida, exist_ok=True)
        
        progress["maximum"] = total_linhas
        progress["value"] = 0
        root.update_idletasks()
        
        for index, row in df.iterrows():
            if cancelar_processo:
                break  # Interrompe o loop se o processo for cancelado
            
            doc = Document()
            
            # Adiciona a logo ao documento (centralizada)
            if os.path.exists(CAMINHO_LOGO):
                para_logo = doc.add_paragraph()
                para_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centraliza a imagem
                run_logo = para_logo.add_run()
                run_logo.add_picture(CAMINHO_LOGO, width=Inches(1.3))  # Ajuste o tamanho conforme necessário
            else:
                messagebox.showwarning("Aviso", f"Logo não encontrada no caminho: {CAMINHO_LOGO}")
                continue  # Continua sem a logo se não for encontrada
            
            # Adiciona o título ao documento
            titulo = doc.add_paragraph()
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = titulo.add_run("<\Feedback Consolidado>")
            run.font.size = TAMANHO_FONTE_TITULO
            run.font.bold = True
            run.font.color.rgb = COR_TEXTO
            run.font.highlight_color = COR_DESTAQUE  # Destaque amarelo escuro
            run.font.name = FONTE
            
            doc.add_paragraph()
            
            # Adiciona os dados da linha ao documento
            nome_aluno = None
            
            for coluna, valor in row.items():
                if cancelar_processo:
                    break  # Interrompe o loop interno se o processo for cancelado
                
                # Verifica se a coluna é "Linha"
                if coluna.strip() == "Linha":
                    # Adiciona um espaço em branco
                    doc.add_paragraph()
                else:
                    if nome_aluno is None:
                        nome_aluno = sanitizar_nome_arquivo(str(valor))  # Sanitiza o nome do aluno
                    
                    paragrafo = doc.add_paragraph()
                    paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinha o texto à esquerda
                    run_coluna = paragrafo.add_run(f"{coluna}: ")
                    run_coluna.font.bold = True
                    run_coluna.font.color.rgb = COR_TEXTO
                    run_coluna.font.highlight_color = COR_DESTAQUE  # Destaque amarelo escuro
                    run_coluna.font.name = FONTE
                    
                    # Substitui valores NaN ou vazios por "Sem informação."
                    if pd.isna(valor) or valor == "":
                        valor_formatado = "Sem informação."
                    else:
                        valor_formatado = str(valor)
                    
                    run_valor = paragrafo.add_run(valor_formatado)
                    run_valor.font.name = FONTE
                    doc.add_paragraph()
            
            # Salva o documento DOCX com o nome do aluno
            if nome_aluno:
                docx_file = os.path.abspath(os.path.join(pasta_saida, f"{nome_aluno}.docx"))
                pdf_file = os.path.abspath(os.path.join(pasta_saida, f"{nome_aluno}.pdf"))
            else:
                docx_file = os.path.abspath(os.path.join(pasta_saida, f"linha_{index+1}.docx"))
                pdf_file = os.path.abspath(os.path.join(pasta_saida, f"linha_{index+1}.pdf"))
            
            doc.save(docx_file)
            
            if not os.path.exists(docx_file):
                messagebox.showerror("Erro", f"Falha ao criar o arquivo DOCX: {docx_file}")
                continue  # Continua para a próxima linha se o DOCX não for criado
            
            # Converte DOCX para PDF
            converter_para_pdf(docx_file, pdf_file)
            
            # Atualiza a barra de progresso
            progress["value"] = index + 1
            root.update()  # Atualiza a interface gráfica para permitir o cancelamento
        
        if not cancelar_processo:
            messagebox.showinfo("Sucesso", MENSAGEM_SUCESSO)
    except Exception as e:
        messagebox.showerror("Erro", MENSAGEM_ERRO.format(e))

# Criar interface gráfica
root = tk.Tk()
root.title("Gerador de PDFs")
root.geometry("500x300")

frame = tk.Frame(root)
frame.pack(pady=20)

label_arquivo = tk.Label(frame, text="Arquivo CSV:")
label_arquivo.grid(row=0, column=0, padx=5, pady=5)
entry_arquivo = tk.Entry(frame, width=40)
entry_arquivo.grid(row=0, column=1, padx=5, pady=5)
btn_arquivo = tk.Button(frame, text="Selecionar", command=selecionar_arquivo)
btn_arquivo.grid(row=0, column=2, padx=5, pady=5)

label_pasta = tk.Label(frame, text="Pasta de Saída:")
label_pasta.grid(row=1, column=0, padx=5, pady=5)
entry_pasta = tk.Entry(frame, width=40)
entry_pasta.grid(row=1, column=1, padx=5, pady=5)
btn_pasta = tk.Button(frame, text="Selecionar", command=selecionar_pasta)
btn_pasta.grid(row=1, column=2, padx=5, pady=5)

progress = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
progress.pack(pady=10)

btn_gerar = tk.Button(root, text="Gerar PDFs", command=gerar_pdfs, width=20)
btn_gerar.pack(pady=10)

btn_cancelar = tk.Button(root, text="Cancelar", command=cancelar, width=20)
btn_cancelar.pack(pady=10)

root.mainloop()